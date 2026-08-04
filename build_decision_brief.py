from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("/Users/zhanglei/Documents/New project/NextPath/NextPath-中考升学助手项目立项决策简报.docx")

FONT_LATIN = "Arial Unicode MS"
FONT_CJK = "Arial Unicode MS"
BLACK = RGBColor(0x00, 0x00, 0x00)
MUTED = RGBColor(0x55, 0x55, 0x55)
BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
LIGHT_GRAY = "F2F4F7"
BORDER = "B7B7B7"
CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run(run, size=11, bold=False, color=BLACK, italic=False):
    run.font.name = FONT_LATIN
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT_LATIN)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_CJK)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tcMar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER, size=4):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = tcPr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(sum(widths_dxa)))
    tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.find(qn("w:tblInd"))
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tblInd.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"), str(widths_dxa[idx]))
            tcW.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[idx] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            set_cell_border(cell)


def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def set_cell_text(cell, text, bold=False, size=9.5, align=WD_ALIGN_PARAGRAPH.LEFT, color=BLACK):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.1
    set_run(p.add_run(str(text)), size=size, bold=bold, color=color)


def add_table(doc, headers, rows, widths, numeric_cols=()):
    table = doc.add_table(rows=1, cols=len(headers))
    set_repeat_table_header(table.rows[0])
    for i, header in enumerate(headers):
        shade_cell(table.rows[0].cells[i], LIGHT_GRAY)
        align = WD_ALIGN_PARAGRAPH.CENTER if i in numeric_cols else WD_ALIGN_PARAGRAPH.LEFT
        set_cell_text(table.rows[0].cells[i], header, bold=True, size=9.5, align=align)
    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data):
            align = WD_ALIGN_PARAGRAPH.CENTER if i in numeric_cols else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(row.cells[i], value, size=9.2, align=align)
    set_table_geometry(table, widths)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    return table


def add_para(doc, text="", bold_prefix=None, italic=False, after=6, keep=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.1
    p.paragraph_format.keep_together = keep
    if bold_prefix and text.startswith(bold_prefix):
        set_run(p.add_run(bold_prefix), bold=True)
        set_run(p.add_run(text[len(bold_prefix):]), italic=italic)
    else:
        set_run(p.add_run(text), italic=italic)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.left_indent = Inches(0.5 if level == 0 else 0.75)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.1
    set_run(p.add_run(text))
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.1
    set_run(p.add_run(text))
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    set_run(p.add_run(text), size={1: 16, 2: 13, 3: 12}[level], bold=True,
            color={1: BLUE, 2: BLUE, 3: DARK_BLUE}[level])
    return p


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT_LATIN)
    fonts.set(qn("w:hAnsi"), FONT_LATIN)
    fonts.set(qn("w:eastAsia"), FONT_CJK)
    rPr.append(fonts)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "20")
    rPr.append(size)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_source(doc, label, url):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.0
    set_run(p.add_run(f"{label}："), size=10, bold=True)
    add_hyperlink(p, url, url)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    set_run(run, size=9, color=MUTED)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = FONT_LATIN
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for level, size, before, after, color in (
        (1, 16, 12, 6, BLUE),
        (2, 13, 10, 5, BLUE),
        (3, 12, 8, 4, DARK_BLUE),
    ):
        style = doc.styles[f"Heading {level}"]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def build():
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run(header.add_run("NextPath｜中考升学助手项目"), size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(footer.add_run("内部讨论稿  ·  "), size=9, color=MUTED)
    add_page_field(footer)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    set_run(p.add_run("立项决策简报"), size=24, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    set_run(p.add_run("NextPath：西安城六区中考升学助手"), size=15, color=MUTED)

    metadata = [
        ("目标", "判断项目是否值得投入，并确定验证方式与预算边界"),
        ("首发市场", "西安城六区，孩子即将进入初三的家长"),
        ("文档版本", "V1.0｜2026年8月3日"),
        ("结论状态", "建议有条件立项，采用分阶段投入和数据验证"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        set_run(p.add_run(f"{label}："), bold=True)
        set_run(p.add_run(value))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    add_heading(doc, "一、决策摘要", 1)
    add_para(doc, "建议有条件立项。该项目解决的是初三家长持续存在且高度刚性的决策问题：孩子当前能上什么层次的高中、距离目标还有多远、学校和政策变化对自家是否有影响，以及中考出分后如何平衡学校层次、管理、出口、距离和录取风险。")
    add_para(doc, "产品不应定义为一次性志愿填报工具，而应定义为初三家庭的中考升学管理助手。它以成绩和排名记录为入口，以动态学校区间、目标差距、个性化提醒、月度规划报告和最终志愿决策为核心交付。")
    add_para(doc, "立项成立的前提：能够取得真实的纵向样本，建立可回测的“初中考试排名→中考区位次→志愿→录取结果”数据链；第一年应以验证模型、用户留存和付费意愿为目标，不宜重资产扩张。")

    add_table(doc,
              ["判断项", "当前结论"],
              [
                  ("用户痛点", "强烈、明确、高频焦虑；现有解决方案割裂且服务体验差"),
                  ("产品空位", "成绩平台与志愿咨询之间存在“持续升学管理”空位"),
                  ("首发市场", "城六区年度核心家庭为5万级，足以验证并形成单城业务"),
                  ("核心壁垒", "本地数据、纵向孩子画像、可校准决策模型和中立信任"),
                  ("主要风险", "数据冷启动、模型过度承诺、单年生命周期和获客成本"),
                  ("资金建议", "首年验证型预算控制在60万至110万元，分阶段拨付"),
              ],
              [1900, 7460])

    add_heading(doc, "二、产品定义", 1)
    add_heading(doc, "2.1 一句话定位", 2)
    add_para(doc, "面向西安城六区初三家长，持续记录孩子成绩和排名，结合本地学校、招生计划和政策，动态判断当前学校层次、目标高中差距和志愿策略，让家长在整个初三阶段心中有数。")

    add_heading(doc, "2.2 核心用户", 2)
    add_bullet(doc, "地域：西安城六区招生单元。")
    add_bullet(doc, "阶段：孩子即将进入初三，从总复习和持续考试阶段开始使用。")
    add_bullet(doc, "决策者：家长，而不是学生。")
    add_bullet(doc, "核心场景：目标明确型家长和位置迷茫型家长。")

    add_heading(doc, "2.3 家长真正购买的答案", 2)
    for item in (
        "孩子现在大概能上什么层次的高中？",
        "想上某所高中，还需提高多少分或前进多少位次？",
        "本次考试是真进步、假进步还是正常波动？",
        "最新招生计划和政策变化与我家是否有关？",
        "如何在学校层次、管理、出口、距离、费用和风险间取得平衡？",
        "哪些学校存在结构性机会，机会依据和风险分别是什么？",
    ):
        add_bullet(doc, item)

    add_heading(doc, "2.4 产品边界", 2)
    add_table(doc,
              ["做什么", "暂不做什么"],
              [
                  ("成绩、排名和目标学校的持续管理", "不在首阶段进入重型错题和整卷分析"),
                  ("大知识点失分标签与提分提醒", "不承诺具体提分结果"),
                  ("个性化学校区间和目标差距", "不把模型包装成确定性算命"),
                  ("相关政策和学校变化提醒", "不做泛资讯信息流"),
                  ("月度升学规划报告", "不以培训班引流作为商业模式"),
                  ("最终志愿方案与风险解释", "不承诺一定录取某校"),
              ],
              [4680, 4680])

    add_heading(doc, "三、核心价值与竞争力", 1)
    add_para(doc, "核心竞争力不是AI、折线图或学校数据库，而是：持续理解一个具体孩子，并把孩子、学校和政策的每一次变化，转化为清晰、可信、个性化的升学判断。")
    add_table(doc,
              ["能力层", "内容", "形成的壁垒"],
              [
                  ("长期孩子画像", "历次成绩、排名、波动、学科结构和目标变化", "临时主播和通用AI无法快速复制"),
                  ("本地升学知识库", "录取位次、招生计划、定向名额、学校管理与出口", "城市化数据持续运营能力"),
                  ("决策模型", "孩子位次分布、学校门槛分布和志愿规则模拟", "可回测、可校准、可解释"),
                  ("信息过滤", "只推送与该家庭相关的变化和行动建议", "从资讯产品升级为决策产品"),
                  ("结果反馈", "中考位次、志愿和真实录取结果反哺模型", "形成逐届增强的数据飞轮"),
                  ("中立信任", "不卖课、不吓唬、不作录取保证", "建立长期品牌资产"),
              ],
              [1600, 3900, 3860])

    add_heading(doc, "四、科学预测与决策依据", 1)
    add_para(doc, "单个孩子最终录取结果只有0或1，但填报前存在可量化的不确定性。科学目标不是猜中一个确定分数，而是得到经过校准的结果分布，并据此选择在多数合理场景下更稳健的方案。")
    add_heading(doc, "4.1 两个基础模型", 2)
    add_number(doc, "孩子最终位次模型：根据所在初中、历次考试、排名百分位、考试类型、趋势和波动，输出中考位次预测区间。")
    add_number(doc, "学校录取门槛模型：根据历年录取位次、招生计划、考生规模、政策和学校热度，输出当年录取门槛区间。")
    add_para(doc, "录取判断来自两个分布的比较，并加入批次、定向资格、志愿顺序和平行志愿投档规则。靠近1至2分边界时，模型应明确标记为高度敏感区，而不是给出虚假精度。")

    add_heading(doc, "4.2 模型验证方式", 2)
    add_bullet(doc, "历史回测：只使用当年填报前可获得的数据预测当年结果。")
    add_bullet(doc, "前瞻验证：中考前冻结预测，中考和录取结束后核对。")
    add_bullet(doc, "关注位次区间覆盖率、学校门槛误差、概率校准、滑档率和方案后悔值。")
    add_bullet(doc, "优先使用可解释的分层贝叶斯、时间状态和情景模拟模型，数据不足时不急于使用黑箱模型。")

    add_heading(doc, "五、数据飞轮与冷启动", 1)
    add_para(doc, "雪球不能从“算法已经很准”开始，而要从“没有私有数据也能提供价值”开始。第一阶段依靠公开数据、确定性规则、产品体验和适度人工服务；第二阶段再用服务过程中沉淀的数据提高模型。")
    add_heading(doc, "5.1 最稀缺的数据链", 2)
    add_para(doc, "所在初中 → 初三历次考试成绩和排名 → 中考成绩与城六区位次 → 志愿顺序 → 最终录取学校。")
    add_heading(doc, "5.2 启动方法", 2)
    add_number(doc, "复盘2026届家庭：收集模考、志愿和录取结果，免费生成全年复盘报告。")
    add_number(doc, "服务2027届种子家庭：持续记录考试，每月形成报告，中考前冻结预测。")
    add_number(doc, "按初中和分数段建立分层模型，样本少的学校借用同类学校规律。")
    add_number(doc, "录取结束后公开模型回测结果和适用边界，逐步建立信任。")

    add_heading(doc, "六、市场与竞品", 1)
    add_heading(doc, "6.1 竞品不是单一软件，而是所有替代方案", 2)
    add_table(doc,
              ["竞品类型", "代表", "优势", "主要空缺"],
              [
                  ("校内成绩平台", "智学网、好分数、七天学堂", "数据丰富、学校入口", "终点是学情和提分，不回答能上什么高中"),
                  ("中考资讯平台", "中考网、中考宝、本地媒体", "资讯和分数线丰富", "信息通用，不了解具体孩子"),
                  ("志愿工具", "各地小程序和临时预测页", "出分后使用简单", "周期短、模型弱、解释不足"),
                  ("本地规划师", "主播、机构和个人顾问", "本地经验和人际信任", "排队、服务不标准、来源不透明"),
                  ("通用AI", "豆包、DeepSeek等", "随时提问、解释能力强", "缺少长期孩子数据和可靠本地规则"),
                  ("高考志愿产品", "掌上高考等", "商业模式已验证", "中考规则高度城市化，不能直接套用"),
              ],
              [1550, 2050, 2700, 3060])
    add_para(doc, "竞争空位：公开市场中尚未发现强势产品将“初三持续成绩管理”和“最终高中及志愿决策”完整连接。机会是建立“中考升学管理”新类别，而不是做另一个成绩查询或志愿查询工具。")

    add_heading(doc, "6.2 市场规模", 2)
    add_para(doc, "全国层面，2024年初中毕业生1698.24万人、普通高中招生1036.20万人。按每个毕业生对应一个家庭决策单元计算，这是一个每年稳定更新的大市场。")
    add_para(doc, "城六区层面，按当前整理的2026年125所高中招生计划计算，总计划约53849人。实际需要升学决策的家庭高于普高招生计划，因此首发市场可以按“至少5万级家庭/年”判断。")
    add_table(doc,
              ["测算场景", "家庭数", "客单价", "年度收入"],
              [
                  ("理论市场（100%）", "53,849", "399元", "约2,149万元"),
                  ("理论市场（100%）", "53,849", "699元", "约3,764万元"),
                  ("付费渗透1%", "约538", "699元", "约38万元"),
                  ("付费渗透5%", "约2,692", "699元", "约188万元"),
                  ("付费渗透10%", "约5,385", "699元", "约376万元"),
              ],
              [2450, 1900, 1900, 3110], numeric_cols=(1, 2, 3))
    add_para(doc, "判断：城六区足以验证产品并形成健康的本地业务，但单城规模有限。长期增长依赖复制到其他大城市，复制对象是“城市数据和模型生产系统”，不是一套全国统一算法。")

    add_heading(doc, "七、商业模式", 1)
    add_table(doc,
              ["层级", "核心交付", "价格假设"],
              [
                  ("免费层", "成绩记录、基础趋势、公开学校信息和有限目标判断", "免费"),
                  ("全年会员", "动态学校区间、目标差距、相关提醒和月度报告", "399至699元/年"),
                  ("决策服务", "中考出分后的志愿方案、情景模拟和人工复核", "1,299至2,999元"),
                  ("未来专业端", "规划师工作台、统一数据和服务标准", "后续验证"),
              ],
              [1800, 5200, 2360])
    add_para(doc, "商业原则：不依赖培训机构佣金和课程导流，避免破坏中立品牌。家庭生命周期约一年，因此必须控制获客成本，并通过同校传播、内容搜索和家长转介绍获得增长。")

    add_heading(doc, "八、品牌与市场推广", 1)
    add_heading(doc, "8.1 品牌定位", 2)
    add_para(doc, "品牌心智：不制造焦虑，让中考升学心中有数。")
    add_para(doc, "核心传播语：孩子现在能上哪，距离目标还有多远。")
    add_para(doc, "品牌原则：不卖课、不训斥、不承诺录取；讲依据、说人话、展示不确定性；所有重要数据标注来源和更新时间。")

    add_heading(doc, "8.2 渠道策略", 2)
    add_table(doc,
              ["渠道", "作用", "建议"],
              [
                  ("微信小程序/服务号", "产品承载、提醒、报告、转介绍", "作为首要阵地"),
                  ("视频号", "本地家长信任和群传播", "数据解读与真实案例"),
                  ("抖音", "规模获客", "不做批评家长式直播，以短结论引流"),
                  ("小红书/搜索", "长尾决策搜索", "学校、位次、通勤和政策专题"),
                  ("家长群/社区", "低成本种子用户和同校传播", "按学校建立种子家庭"),
                  ("线下规划师", "高复杂案例和专业复核", "后期接入，先不做平台"),
              ],
              [2100, 3000, 4260])

    add_heading(doc, "8.3 内容策略", 2)
    add_bullet(doc, "一次考试怎么解读：涨分但排名不变意味着什么。")
    add_bullet(doc, "一所学校讲透：层次、管理、出口、距离和录取风险。")
    add_bullet(doc, "一条政策只讲影响谁：与哪些位次、初中和目标家庭相关。")
    add_bullet(doc, "一个真实家庭复盘：判断依据、结果和模型偏差。")
    add_bullet(doc, "一次公开回测：预测在哪些场景可靠、在哪些场景不可靠。")

    add_heading(doc, "九、投入成本核算", 1)
    add_para(doc, "以下为首年现金投入的规划假设，不包含大规模全国扩张。成本会受创始人是否亲自承担产品和研发、是否使用外包、人工审核比例和获客方式影响。")

    add_heading(doc, "9.1 三种投入方式", 2)
    add_table(doc,
              ["方案", "团队形态", "首年投入", "适用判断"],
              [
                  ("验证型（建议）", "创始人主导；1名技术核心；兼职设计、数据和运营", "60万至110万元", "先验证数据、留存、模型和付费"),
                  ("标准型", "4至5人全职产品技术运营团队", "150万至250万元", "已有明确需求和稳定获客渠道"),
                  ("进攻型", "7至9人团队并投入内容和投放", "300万至500万元", "模型与转化已验证后扩城"),
              ],
              [1800, 3500, 1800, 2260])

    add_heading(doc, "9.2 建议采用的首年验证型预算", 2)
    add_table(doc,
              ["成本项", "预算范围", "说明"],
              [
                  ("人员与技术开发", "35万至60万元", "技术核心、兼职设计和必要外包；创始人承担产品与部分研发可下降"),
                  ("数据整理与运营", "8万至15万元", "学校、政策、计划、历史位次和家庭样本清洗核验"),
                  ("云服务、模型与工具", "2万至6万元", "服务器、数据库、消息、地图、模型API和监控"),
                  ("种子用户与研究", "3万至6万元", "家庭复盘、激励、访谈和测试"),
                  ("品牌内容与推广", "8万至20万元", "内容生产、账号运营和小规模投放"),
                  ("法律、隐私与合规", "3万至8万元", "协议、儿童数据规则、安全和广告表述审核"),
                  ("其他及预备金", "6万至10万元", "设备、差旅、不可预见支出"),
              ],
              [2500, 2000, 4860])
    add_para(doc, "建议预算上限：首年不超过110万元；首期只批准15万至20万元，用于数据验证、产品定义、2026届复盘和2027届种子家庭启动。达到阶段指标后再拨付后续预算。")

    add_heading(doc, "9.3 收入与盈亏平衡敏感性", 2)
    add_table(doc,
              ["付费家庭", "全年会员客单价", "会员收入", "判断"],
              [
                  ("500户", "699元", "约35万元", "不能覆盖验证型首年投入"),
                  ("1,500户", "699元", "约105万元", "接近验证型现金投入上限"),
                  ("3,000户", "699元", "约210万元", "具备单城正向经营可能"),
                  ("1,500户", "899元混合客单价", "约135万元", "需要会员与志愿服务组合"),
              ],
              [2000, 2350, 2200, 2810], numeric_cols=(0, 1, 2))
    add_para(doc, "粗略盈亏平衡判断：若首年投入80万元、自动化服务的单户贡献毛利约500至650元，需要约1,230至1,600个付费家庭。若大量使用人工一对一服务，所需付费家庭数会显著提高。")

    add_heading(doc, "十、主要风险与应对", 1)
    add_table(doc,
              ["风险", "表现", "应对原则"],
              [
                  ("数据冷启动", "缺少平时排名到中考位次的配对样本", "先用公开数据和宽区间提供价值，同时做2026复盘与2027前瞻队列"),
                  ("虚假精确", "边界1至2分却输出确定结论", "输出区间、敏感度和置信等级，公开回测"),
                  ("数据偏差", "样本集中在重点初中和高分家庭", "主动覆盖不同学校、区域和分数段"),
                  ("获客成本", "家庭只使用一年，付费投放难回收", "微信生态、同校转介绍和搜索内容优先"),
                  ("学校主观信息", "管理和出口信息混入传闻", "建立官方、交叉验证和单一反馈的来源分级"),
                  ("大平台进入", "成绩平台向升学决策延伸", "建立家长侧跨平台档案、本地模型和中立品牌"),
                  ("合规风险", "儿童数据、AI输出和升学承诺", "监护人同意、最小采集、可删除、广告和模型表述审核"),
              ],
              [1900, 3200, 4260])

    add_heading(doc, "十一、建议的立项方式与决策门槛", 1)
    add_para(doc, "建议不是一次性批准完整项目，而是采用阶段门制度。项目只有在用户价值、数据可得性、模型可验证性和获客效率同时成立时才继续投入。")
    add_table(doc,
              ["阶段", "核心任务", "继续投入的门槛"],
              [
                  ("阶段一：立项验证", "2026届复盘、用户访谈、数据字典、产品方案和首批种子家庭", "能取得真实完整案例；家长明确认可核心价值；数据输入可持续"),
                  ("阶段二：年度陪伴", "服务2027届家庭，持续月报、目标差距和相关提醒", "考试记录留存、报告使用和付费意愿达到预设目标"),
                  ("阶段三：中考验证", "冻结预测、志愿方案、结果回收和模型回测", "预测区间覆盖和志愿结果明显优于简单使用去年分数线"),
                  ("阶段四：商业扩张", "扩大城六区渗透，并验证第二个城市", "获客成本、贡献毛利和城市数据复制成本可控"),
              ],
              [1800, 4000, 3560])

    add_heading(doc, "十二、最终建议", 1)
    add_para(doc, "项目值得做，但应按“数据和服务驱动的本地决策产品”立项，而不是按普通教育App或一次性志愿工具立项。第一年目标不是追求用户规模，而是证明四件事：家长愿意持续记录；月报和目标差距能够形成留存；纵向数据可以支持比简单规则更好的判断；家长愿意为全年清晰决策付费。")
    add_para(doc, "建议批准小规模验证，并设置首年110万元绝对预算上限。最先投入的不是大规模开发和广告，而是数据标准、种子家庭、真实案例复盘、产品体验和模型回测体系。")

    add_heading(doc, "附录：关键数据与公开来源", 1)
    add_source(doc, "教育部《2024年全国教育事业发展统计公报》", "https://www.moe.gov.cn/jyb_sjzl/sjzl_fztjgb/202506/t20250611_1193760.html")
    add_source(doc, "西安市2026年高中阶段学校招生录取工作方案", "https://edu.xa.gov.cn/xwzx/tzgg/2076480238905311234.html")
    add_source(doc, "西安市2026年普通高中分学校招生计划", "https://edu.xa.gov.cn/xwzx/tzgg/2077576360413937666.html")
    add_source(doc, "西安市2026年城六区普通高中录取控制线", "https://edu.xa.gov.cn/xwzx/tzgg/2077952493749297153.html")
    add_source(doc, "《生成式人工智能服务管理暂行办法》", "https://www.gov.cn/zhengce/zhengceku/202307/content_6891752.htm")
    add_source(doc, "《儿童个人信息网络保护规定》", "https://www.gov.cn/zhengce/zhengceku/2019-08/22/content_5458118.htm")
    add_source(doc, "《中华人民共和国广告法》", "https://www.samr.gov.cn/zw/zfxxgk/fdzdgknr/fgs/art/2023/art_5474cf75173c45d6a0379730fb4e8d97.html")
    add_source(doc, "概率预测的校准与锐度研究", "https://doi.org/10.1111/j.1467-9868.2007.00587.x")
    add_source(doc, "Distribution-Free Predictive Inference for Regression", "https://doi.org/10.1080/01621459.2017.1307116")

    core = doc.core_properties
    core.title = "NextPath中考升学助手项目立项决策简报"
    core.subject = "内部立项讨论"
    core.author = "NextPath项目组"
    core.keywords = "中考,升学规划,产品立项,市场分析,成本核算"
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
